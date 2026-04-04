"""
=============================================================
  CẤU TRÚC FILE KẾ TOÁN - STRUCTURAL PARSER  (Cách 2)
  Hỗ trợ: .txt (đã convert từ PDF scan)
           .pdf (text layer)
           .pdf (scan / ảnh) → OCR
           .docx
  Output: dict chuẩn JSONL cho PhoBERT / mPhoBERT
=============================================================
"""

import re
import json
import unicodedata
from pathlib import Path

# PHẦN 1 — TIỆN ÍCH DÙNG CHUNG

def normalize(text: str) -> str:
    """Chuẩn hóa Unicode NFC (bắt buộc với tiếng Việt)."""
    return unicodedata.normalize("NFC", text)

def remove_ocr_noise(text: str) -> str:
    """
    Làm sạch text đã qua OCR từ TXT. (Hỗ trợ nhiễu nặng)
    File BVTA có ký tự rác: oe, l thay 1, ký tự ASCII lạ...
    """
    # Xóa ký tự điều khiển (giữ newline)
    text = re.sub(r"[^\S\n]+", " ", text)
    # Xóa gạch ngang/chấm/underscore kéo dài (lỗi vỡ bảng)
    text = re.sub(r"[_\-\.]{5,}", " ", text)
    # Xóa dòng chỉ chứa toàn ký tự đặc biệt hoặc rác không có chữ/số
    text = re.sub(r"^\s*[^a-zA-Z0-9À-ỹ\n]{5,}\s*$", "", text, flags=re.MULTILINE)
    # Xóa header/footer trang lặp lại (ví dụ "- 5 -", "Page 5", "Trang 1/2")
    text = re.sub(r"^\s*(?:[-–—]?\s*\d+\s*[-–—]?|Trang\s*\d+(?:/\d+)?|Page\s*\d+)\s*$", "", text, flags=re.MULTILINE|re.IGNORECASE)
    # Gộp nhiều dòng trống thành 1
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def clean_number(raw: str) -> float | None:
    """Chuyển chuỗi số kế toán VN (1.234.567, 1,234,567, 1 234 567) → float."""
    if not raw:
        return None
    # Thay thế các khoảng trắng (bao gồm cả OCR nhận nhầm dấu cách ở giữa số)
    s = re.sub(r"\s+", "", str(raw))
    # Xử lý số âm trong ngoặc hoặc dấu trừ: (123.456) hoặc -123.456
    negative = s.startswith("(") and s.endswith(")") or s.startswith("-")
    s = s.strip("()-")
    # Loại bỏ dấu phân cách hàng nghìn (dấu phẩy hoặc chấm đứng trước 3 chữ số liên tiếp)
    s = re.sub(r"[,\.](?=\d{3}(?:[,\.]|\b))", "", s)
    # Nếu còn dấu ',' hoặc '.' ở cuối → dấu thập phân
    s = s.replace(",", ".")
    try:
        val = float(s)
        return -val if negative else val
    except ValueError:
        return None

# PHẦN 2 — PHÁT HIỆN LOẠI FILE VÀ ĐỌC TEXT THÔ

def read_pdf_text_layer(path: Path) -> tuple[str, list[dict]]:
    """
    Trích xuất Text bằng PyMuPDF (fitz) - Hiệu năng cao nhất cho Báo cáo tài chính
    theo khuyến nghị từ bài báo arXiv:2410.09871v2.
    """
    import fitz  # PyMuPDF

    pages_text = []
    tables_extracted = []

    with fitz.open(path) as doc:
        for page_num, page in enumerate(doc, start=1):
            # ── 1. Trích xuất Text giữ nguyên Layout ───────────
            # Tham số 'layout' để giữ khoảng cách cột giống báo cáo
            text = page.get_text("layout") or ""
            pages_text.append(text)

            # ── 2. Trích xuất Bảng (Cơ bản bằng PyMuPDF) ───────
            # PyMuPDF >= 1.23 hỗ trợ page.find_tables()
            try:
                tabs = page.find_tables()
                if tabs and tabs.tables:
                    for tbl in tabs.tables:
                        clean_rows = []
                        for row in tbl.extract():
                            clean_row = [str(cell or "").strip() for cell in row]
                            if any(clean_row):           # bỏ hàng trống hoàn toàn
                                clean_rows.append(clean_row)
                        if clean_rows:
                            tables_extracted.append({
                                "page": page_num,
                                "rows": clean_rows,
                            })
            except Exception:
                pass  # Fallback nếu PyMuPDF version cũ

    return "\n\n".join(pages_text), tables_extracted


def extract_table_with_tatr(path: Path):
    """
    [Integration Hook] - TATR (Table Transformer)
    Khuyến nghị từ arXiv:2410.09871v2 cho Financial Table Detection.
    Yêu cầu cài đặt: pip install transformers timm torchvision
    Sẽ được kích hoạt khi xử lý các bảng Thuyết minh / Ngoại bảng siêu phức tạp.
    """
    pass


def read_docx(path: Path) -> tuple[str, list[dict]]:
    """
    File Word (.docx) → dùng python-docx.
    Trả về: (full_text, danh_sach_bang)
    Giữ nguyên cấu trúc Heading / Paragraph / Table.
    """
    from docx import Document

    doc = Document(path)
    parts = []
    tables_extracted = []

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]   # lấy phần tên tag bỏ namespace

        if tag == "p":
            # Đoạn văn thông thường hoặc Heading
            para = None
            for p in doc.paragraphs:
                if p._element is block:
                    para = p
                    break
            if para and para.text.strip():
                style = para.style.name   # "Heading 1", "Normal", ...
                prefix = ""
                if "Heading" in style:
                    level = re.search(r"\d+", style)
                    prefix = "#" * int(level.group()) + " " if level else "# "
                parts.append(prefix + para.text.strip())

        elif tag == "tbl":
            # Bảng biểu trong DOCX
            tbl = None
            for t in doc.tables:
                if t._element is block:
                    tbl = t
                    break
            if tbl:
                rows = []
                for row in tbl.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        rows.append(row_data)
                if rows:
                    tables_extracted.append({"rows": rows})
                    # Thêm text đại diện của bảng vào full_text
                    for row in rows:
                        parts.append(" | ".join(row))

    return "\n".join(parts), tables_extracted



# PHẦN 3 — PHÂN TÍCH CẤU TRÚC BÁO CÁO TÀI CHÍNH

# Pattern nhận biết từng loại báo cáo trong BCTC Việt Nam
SECTION_PATTERNS = {
    "bang_cdkt": re.compile(
        r"(B[AÀ]NG?\s*C[AÂ]N?\s*[DĐ][OÔ]I?\s*K[EÊ]?\s*TO[AÁ]N"
        r"|BALANCE\s*SHEET|B\s*01"
        r"|BANG\s*CAN\s*DOI\s*KE\s*TOAN)",    # không dấu từ OCR
        re.I
    ),
    "kqkd": re.compile(
        r"(K[EÊ]T?\s*QU[AÁ]?\s*HO[AẠ]T?\s*[DĐ][OÔ]NG"
        r"|KQ\s*H[DĐ]KD|INCOME\s*STATEMENT|B\s*02"
        r"|KET\s*QUA\s*HOAT\s*DONG)",          # không dấu từ OCR
        re.I
    ),
    "lctt": re.compile(
        r"(L[UƯ]U?\s*CHUY[EÊ]N?\s*TI[EÊ]N?\s*T[EÊ]"
        r"|CASH\s*FLOW|B\s*03"
        r"|LUU\s*CHUYEN\s*TIEN\s*TE)",         # không dấu từ OCR
        re.I
    ),
    "thuyet_minh": re.compile(
        r"(THUY[EÊ]T?\s*MINH|NOTES?\s*TO\s*THE|B\s*09"
        r"|THUYET\s*MINH)",                     # không dấu từ OCR
        re.I
    ),
}

def detect_sections(text: str) -> dict[str, str]:
    """
    Tách BCTC thành 4 phần chính:
    bang_cdkt, kqkd, lctt, thuyet_minh
    Trả về dict {section_name: nội_dung}
    """
    # Tìm vị trí bắt đầu mỗi section
    positions = {}
    for section_name, pattern in SECTION_PATTERNS.items():
        m = pattern.search(text)
        if m:
            positions[section_name] = m.start()

    if not positions:
        return {"full_text": text}

    # Sắp xếp theo thứ tự xuất hiện
    sorted_sections = sorted(positions.items(), key=lambda x: x[1])
    sections = {}

    for i, (name, start) in enumerate(sorted_sections):
        end = sorted_sections[i + 1][1] if i + 1 < len(sorted_sections) else len(text)
        sections[name] = text[start:end].strip()

    return sections


def split_txt_by_page(text: str) -> list[dict]:
    """
    TXT file có dạng:
        --- TRANG 1 ---
        ...nội dung...
        --- TRANG 2 ---
        ...
    → Tách thành list[{"page": int, "content": str}]
    """
    pattern = re.compile(r"---\s*TRANG\s*(\d+)\s*---", re.I)
    parts = pattern.split(text)

    pages = []
    # parts = [text_before_first, page_num, content, page_num, content, ...]
    i = 1
    while i < len(parts) - 1:
        page_num = int(parts[i])
        content = parts[i + 1].strip()
        pages.append({"page": page_num, "content": content})
        i += 2

    # Nếu không có marker trang → trả về toàn bộ như 1 trang
    if not pages:
        pages = [{"page": 1, "content": text.strip()}]

    return pages

# PHẦN 4 — TRÍCH XUẤT FIELD KẾ TOÁN BẰNG REGEX

# Field patterns — Cải tiến khoảng cách và chống nhiễu OCR
FIELD_PATTERNS = {
    "ten_cong_ty": [
        # Tiếng Việt - lấy sau từ khóa công ty, độ dài mở rộng
        r"(?:^|\n)(?:Công\s*ty|CÔNG\s*TY)\s+(?:CP|TNHH|Cổ\s*Phần|C[OÔ]NG\s*TY)?\s*(.{5,120}?)(?:\n|MST|Mã\s*số|Địa|$)",
        # Tiếng Anh
        r"^([A-Z][A-Z\s&,\.]{10,120}(?:COMPANY|JSC|CORPORATION|CO\.,?\s*LTD))",
    ],
    "ky_bao_cao": [
        r"(?:Quý|QUÍ|Quarter|Q)\s*(\d)\s*(?:Năm|NĂM|Year|NAM)?\s*(\d{4})",
        r"(?:(?:N|n)ăm\s*tài\s*chính|năm|year|NAM)\s+(\d{4})",
        r"(?:Kết\s*thúc\s*ngày.*?|Cho\s*kỳ\s*kế.*?)(\d{4})",
    ],

    # ── Các khoản mục giá trị tiền: cho phép khoảng cách lớn để nhảy qua Text nhiễu, Note/Thuyết minh ──
    "doanh_thu": [
        # VN: "Doanh thu" cho phép cách chữ số đến 120 ký tự (không phải số) để tránh rác
        r"[Dd]oanh\s*thu\s*(?:thu[ầa]n|b[áa]n\s*h[àa]ng|ho[ạa]t\s*đ[ộo]ng).*?[^\d]{0,120}([\d]{1,3}(?:[\s,\.]+\d{3}){2,})",
        # EN: Net revenue
        r"Net\s*revenue.*?[^\d]{0,120}([\d]{1,3}(?:[\s,\.]+\d{3}){2,})",
        r"Revenue\s*from\s*sales.*?[^\d]{0,120}([\d]{1,3}(?:[\s,\.]+\d{3}){2,})",
    ],
    "loi_nhuan_truoc_thue": [
        # VN — mã khoản 50
        r"L[oợ]i\s*nhu[ậa]n\s*(?:k[ếê]\s*to[áa]n\s*)?tr[ướu][ớốc]c?\s*thu[ếê].*?[^\d]{0,120}(\(?[\d]{1,3}(?:[\s,\.]+\d{3}){2,}\)?)",
        r"Accounting\s*profit\s*before\s*tax.*?[^\d]{0,120}(\(?[\d]{1,3}(?:[\s,\.]+\d{3}){2,}\)?)",
        r"\b(?:Mã\s*số\s*)?50\b.*?[^\d]{0,80}([\d]{1,3}(?:[\s,\.]+\d{3}){2,})",
    ],
    "loi_nhuan_sau_thue": [
        # LNST cho phép số âm định dạng kế toán (...)
        r"L[oợ]i\s*nhu[ậa]n\s*sau\s*thu[ếê][^\d\(\)]*?[^\d\(\)]{0,120}(\(?[\d]{1,3}(?:[\s,\.]+\d{3}){2,}\)?)",
        r"Profit\s*after\s*tax.*?[^\d]{0,120}(\(?[\d]{1,3}(?:[\s,\.]+\d{3}){2,}\)?)",
        r"\b(?:Mã\s*số\s*)?60\b.*?[^\d]{0,80}(\(?[\d]{1,3}(?:[\s,\.]+\d{3}){2,}\)?)",
    ],
    "tong_tai_san": [
        r"(?:T[ổÔo]NG\s*(?:C[ỘÔo]NG\s*)?)?T[ÀA]I\s*S[ẢA]N.*?[^\d]{0,120}([\d]{1,3}(?:[\s,\.]+\d{3}){2,})",
        r"TOTAL\s*ASSETS.*?[^\d]{0,120}([\d]{1,3}(?:[\s,\.]+\d{3}){2,})",
        r"\b(?:Mã\s*số\s*)?270\b.*?[^\d]{0,80}([\d]{1,3}(?:[\s,\.]+\d{3}){2,})",
    ],
    "von_chu_so_huu": [
        r"V[ỐÔo]N\s*CH[ỦU]\s*S[ởỞo]\s*H[ỮÙU]U?.*?[^\d]{0,120}([\d]{1,3}(?:[\s,\.]+\d{3}){2,})",
        r"EQUITY.*?[^\d]{0,120}([\d]{1,3}(?:[\s,\.]+\d{3}){2,})",
        r"\b(?:Mã\s*số\s*)?400\b.*?[^\d]{0,80}([\d]{1,3}(?:[\s,\.]+\d{3}){2,})",
    ],
    "tien_va_tuong_duong": [
        r"Ti[eề]n\s*v[àa]\s*c[áa]c\s*kho[ảa]n\s*t[ưươo][oơ]ng\s*[đd][ưươo][ơo]ng.*?[^\d]{0,120}([\d]{1,3}(?:[\s,\.]+\d{3}){2,})",
        r"Cash\s*and\s*cash\s*equivalents.*?[^\d]{0,120}([\d]{1,3}(?:[\s,\.]+\d{3}){2,})",
        r"\b(?:Mã\s*số\s*)?110\b.*?[^\d]{0,80}([\d]{1,3}(?:[\s,\.]+\d{3}){2,})",
    ],
}

def extract_fields(text: str) -> dict:
    """
    Trích xuất các field tài chính chính từ text.
    Trả về dict với giá trị đã chuẩn hóa.
    """
    fields = {}
    for field, patterns in FIELD_PATTERNS.items():
        for pat in patterns:
            m = re.search(pat, text, re.IGNORECASE | re.MULTILINE)
            if m:
                raw = " ".join(g for g in m.groups() if g).strip()
                # Thử chuyển về số nếu là số tiền
                num = clean_number(raw)
                fields[field] = num if num is not None else raw
                break
        if field not in fields:
            fields[field] = None

    return fields



# PHẦN 5 — PHÁT HIỆN DẤU HIỆU GIAN LẬN (RULE-BASED)

def detect_fraud_signals(text: str, fields: dict) -> list[str]:
    """
    Phát hiện dấu hiệu gian lận theo luật kế toán VN.
    Trả về danh sách các dấu hiệu tìm thấy.
    """
    signals = []

    # 1. Lợi nhuận âm bất thường (lỗ nặng)
    lnst = fields.get("loi_nhuan_sau_thue")
    doanh_thu = fields.get("doanh_thu")
    if isinstance(lnst, (int, float)) and isinstance(doanh_thu, (int, float)):
        if doanh_thu > 0 and lnst < 0:
            ratio = abs(lnst) / doanh_thu
            if ratio > 0.3:
                signals.append(f"lo_nang_{ratio:.1%}_doanh_thu")

    # 2. Lỗ nhiều quý liên tiếp (phát hiện qua text)
    if re.search(r"lỗ.{0,20}(?:liên\s*tiếp|nhiều\s*quý|kéo\s*dài)", text, re.I):
        signals.append("lo_lien_tiep")

    # 3. Số tiền tròn bất thường (Benford's Law heuristic)
    amounts = re.findall(r"\b(\d+)(?:,000){2,}\b", text)
    round_count = sum(1 for a in amounts if a.endswith("000"))
    if round_count > 5:
        signals.append(f"so_tien_tron_bat_thuong_{round_count}_lan")

    # 4. Từ khóa nghi vấn
    fraud_keywords = [
        r"hoa\s*hồng\s*(?:lớn|cao|bất\s*thường)",
        r"chi\s*phí\s*không\s*rõ\s*nguồn\s*gốc",
        r"quà\s*biếu",
        r"chi\s*phí\s*tiếp\s*khách\s*(?:lớn|bất\s*thường)",
        r"related\s*party\s*transaction",        # tiếng Anh
        r"(?:không|chưa)\s*có\s*chứng\s*từ",
    ]
    for kw in fraud_keywords:
        if re.search(kw, text, re.I):
            signals.append(f"tu_khoa_{kw[:20].strip()}")

    # 5. Dự phòng phải thu khó đòi quá lớn
    if re.search(r"dự\s*phòng.{0,50}(?:[\d\.]{8,})", text, re.I):
        m = re.search(r"dự\s*phòng.{0,50}([\d\.]{8,})", text, re.I)
        if m:
            prov = clean_number(m.group(1))
            if prov and isinstance(doanh_thu, (int, float)) and doanh_thu > 0:
                if prov / doanh_thu > 0.1:
                    signals.append("du_phong_lon_hon_10pct_dt")

    # 6. Khoản phải thu liên quan bên thứ ba chiếm tỷ trọng lớn
    if re.search(r"related\s*party|bên\s*liên\s*quan", text, re.I):
        signals.append("co_giao_dich_ben_lien_quan")

    return signals



# PHẦN 6 — HÀM CHÍNH: XỬ LÝ 1 FILE

def process_file(path: Path) -> dict | None:
    """
    Xử lý 1 file bất kỳ (.txt / .pdf / .docx).
    Trả về dict chuẩn sẵn sàng ghi JSONL.
    """
    ext = path.suffix.lower()
    result = {
        "file":        path.name,
        "source_type": None,
        "raw_pages":   [],
        "sections":    {},
        "tables":      [],
        "fields":      {},
        "fraud_signals": [],
        "label":       -1,            # -1 = chưa gắn nhãn
        "suggested_review": False,
    }

    try:
        # ── Đọc theo loại file ───────────────────────────────
        if ext == ".txt":
            raw = read_txt(path)
            raw = normalize(raw)
            raw = remove_ocr_noise(raw)
            result["source_type"] = "txt_converted"

            # Tách theo trang (--- TRANG X ---)
            result["raw_pages"] = split_txt_by_page(raw)
            full_text = raw

        elif ext == ".pdf":
            # Trích xuất 100% bằng PyMuPDF trên Text Layer gốc (Tắt hoàn toàn OCR)
            full_text, tables = read_pdf_text_layer(path)
            full_text = normalize(full_text)
            full_text = remove_ocr_noise(full_text)
            result["source_type"] = "pdf_text_layer"
            result["tables"] = tables
            result["raw_pages"] = [
                {"page": i + 1, "content": p}
                for i, p in enumerate(full_text.split("\n\n"))
            ]

        elif ext in (".docx", ".doc"):
            full_text, tables = read_docx(path)
            full_text = normalize(full_text)
            result["source_type"] = "docx"
            result["tables"] = tables
            result["raw_pages"] = [{"page": 1, "content": full_text}]

        else:
            return None   # Loại file không hỗ trợ

        # ── Phân tích cấu trúc báo cáo ───────────────────────
        result["sections"] = detect_sections(full_text)

        # ── Trích field tài chính ─────────────────────────────
        result["fields"] = extract_fields(full_text)

        # ── Phát hiện dấu hiệu gian lận ──────────────────────
        result["fraud_signals"] = detect_fraud_signals(
            full_text, result["fields"]
        )
        result["suggested_review"] = len(result["fraud_signals"]) > 0

        # Lưu snippet text sạch (không lưu toàn bộ để tiết kiệm bộ nhớ)
        result["text_snippet"] = full_text[:500]
        result["char_count"] = len(full_text)

        return result

    except Exception as e:
        result["error"] = str(e)
        result["source_type"] = "error"
        return result


# PHẦN 7 — CHẠY HÀNG LOẠT VÀ GHI DATASET

def run_pipeline(
    input_dirs: list[str],
    output_path: str = "dataset.jsonl",
    max_workers: int = 4,
):
    from tqdm import tqdm
    from concurrent.futures import ProcessPoolExecutor, as_completed

    all_files = []
    for d in input_dirs:
        folder = Path(d)
        for ext in ("*.txt", "*.pdf", "*.docx"):
            all_files.extend(folder.rglob(ext))

    print(f"\nTìm thấy {len(all_files)} file — bắt đầu xử lý...\n")

    results, errors = [], []

    # ProcessPoolExecutor để chạy song song
    with ProcessPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(process_file, f): f for f in all_files}
        for future in tqdm(as_completed(futures), total=len(all_files), desc="Processing"):
            res = future.result()
            if res is None:
                continue
            if "error" in res:
                errors.append(res)
            else:
                results.append(res)

    # Ghi dataset chính
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    with out.open("w", encoding="utf-8") as f:
        for r in results:
            f.write(json.dumps(r, ensure_ascii=False) + "\n")

    # Ghi log lỗi riêng
    if errors:
        err_path = out.with_name("errors.jsonl")
        with err_path.open("w", encoding="utf-8") as f:
            for e in errors:
                f.write(json.dumps(e, ensure_ascii=False) + "\n")

    # Thống kê
    review_needed = sum(1 for r in results if r.get("suggested_review"))
    src_types     = {}
    for r in results:
        t = r.get("source_type", "unknown")
        src_types[t] = src_types.get(t, 0) + 1

    print(f"\n{'='*50}")
    print(f"  Đã xử lý : {len(results):,} file")
    print(f"  Lỗi      : {len(errors):,} file  (xem errors.jsonl)")
    print(f"  Cần xem  : {review_needed:,} file có dấu hiệu gian lận")
    print(f"  Loại nguồn: {src_types}")
    print(f"  Dataset  : {output_path}")
    print(f"{'='*50}")
